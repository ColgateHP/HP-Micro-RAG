import os
import io
import re
import json
from typing import List, Tuple
import fitz  # PyMuPDF
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from langchain.docstore.document import Document as LangchainDocument

def get_gdrive_service():
    """初始化并返回Google Drive API服务实例。"""
    creds_json = os.getenv('GCP_SA_KEY', '{}')
    creds_info = json.loads(creds_json)
    creds = service_account.Credentials.from_service_account_info(
        creds_info, scopes=['https://www.googleapis.com/auth/drive']
    )
    return build('drive', 'v3', credentials=creds)

def _get_id_from_url(url: str) -> str:
    """从Google Drive URL中提取文件ID。"""
    match = re.search(r'/d/([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)
    raise ValueError(f"无法从URL中提取ID: {url}")

def _download_and_parse_file(service, file_id: str, doc_name: str) -> LangchainDocument | None:
    """根据文件ID下载、解析单个文件。"""
    try:
        file_metadata = service.files().get(fileId=file_id, fields='mimeType').execute()
        mime_type = file_metadata.get('mimeType')

        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)

        content = ""
        if 'application/pdf' in mime_type:
            with fitz.open(stream=fh, filetype="pdf") as pdf_doc:
                content = "".join(page.get_text() for page in pdf_doc)
        elif 'application/vnd.google-apps.document' in mime_type:
            # 对于Google Docs, 需要导出为docx格式
            export_request = service.files().export_media(fileId=file_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            fh_export = io.BytesIO()
            downloader_export = MediaIoBaseDownload(fh_export, export_request)
            done_export = False
            while not done_export:
                _, done_export = downloader_export.next_chunk()
            fh_export.seek(0)
            doc = Document(fh_export)
            content = "\n".join(para.text for para in doc.paragraphs)

        if content:
            return LangchainDocument(page_content=content, metadata={"source": doc_name})
        else:
            print(f"[WARNING] 文件内容为空或不支持的格式: {doc_name} (MimeType: {mime_type})")
            return None

    except HttpError as e:
        print(f"[ERROR] 下载文件 {doc_name} (ID: {file_id}) 时发生HTTP错误: {e}")
        return None
    except Exception as e:
        print(f"[ERROR] 处理文件 {doc_name} 时发生未知错误: {e}")
        return None


def load_documents_from_txt(file_path: str) -> List[LangchainDocument]:
    """
    读取txt文件，解析每一行的'名称: URL'，并加载所有文档。
    """
    if not os.path.exists(file_path):
        print(f"[ERROR] 输入文件未找到: {file_path}")
        return []

    service = get_gdrive_service()
    docs = []
    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or ':' not in line:
                continue
            
            doc_name, url = line.split(':', 1)
            doc_name = doc_name.strip()
            url = url.strip()
            
            try:
                print(f"[INFO] 开始处理文档: {doc_name}")
                file_id = _get_id_from_url(url)
                doc = _download_and_parse_file(service, file_id, doc_name)
                if doc:
                    docs.append(doc)
            except ValueError as e:
                print(e)
    
    return docs
